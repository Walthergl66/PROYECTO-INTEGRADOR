"""Genera el documento Word de la Fase 1."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[2]
INPUT = RAIZ / "informe" / "FASE_1_DEFINICION_NEGOCIO.md"
OUTPUT = RAIZ / "informe" / "FASE_1_DEFINICION_NEGOCIO.docx"


def set_run(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph_with_bold(doc: Document, text: str, style: str | None = None):
    paragraph = doc.add_paragraph(style=style)
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        set_run(run, bold=(i % 2 == 1))
    return paragraph


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SafeAnalytics EC")
    set_run(run, size=22, bold=True, color="12355B")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Sistema Inteligente de Analítica de Negocios para el Monitoreo y "
        "Análisis Estratégico de Homicidios Intencionales en Ecuador"
    )
    set_run(run, size=12, color="17202A")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Proyecto Integrador - Analítica de Negocios 7A")
    set_run(run, size=11, bold=True, color="0F766E")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Fase 1: Definición del Negocio y Entendimiento del Problema")
    set_run(run, size=15, bold=True, color="2E74B5")
    doc.add_page_break()


def build_doc() -> None:
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    for raw_line in INPUT.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            item = line[2:]
            p = doc.add_paragraph(style="List Bullet")
            parts = item.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                set_run(run, bold=(i % 2 == 1))
        else:
            add_paragraph_with_bold(doc, line)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("SafeAnalytics EC - Fase 1")
    set_run(run, size=9, color="617080")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
